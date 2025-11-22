from typing import List, Dict, Any
import os

import logging
import re
import hashlib
import json
from datetime import datetime
from pathlib import Path
import pymupdf  # PyMuPDF for PDF processing
import pandas as pd
from langchain.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import Chroma
from langchain.embeddings import OpenAIEmbeddings
from langchain.schema import Document
from app.services.database import DatabaseService
try:
    from docx import Document as DocxDocument
    from docx.table import Table as DocxTable
    from docx.text.paragraph import Paragraph as DocxParagraph
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    logging.warning("DOCX processing libraries not available. Install python-docx for Word document support.")

try:
    from arabic_reshaper import reshape
    from bidi.algorithm import get_display
    PERSIAN_SUPPORT = True
except ImportError:
    PERSIAN_SUPPORT = False
    logging.warning("Persian text processing libraries not available. Install arabic-reshaper and python-bidi for full Persian support.")

from app.core.config import settings

class OpenRouterEmbeddings:
    def __init__(self, api_key: str, base_url: str, model: str, referer: str = None, site_title: str = None):
        from openai import OpenAI
        self.client = OpenAI(base_url=base_url, api_key=api_key)
        self.model = model
        self.extra_headers = {}
        if referer:
            self.extra_headers["HTTP-Referer"] = referer
        if site_title:
            self.extra_headers["X-Title"] = site_title

    def embed_query(self, text: str):
        resp = self.client.embeddings.create(
            model=self.model,
            input=text,
            encoding_format="float",
            extra_headers=self.extra_headers or None,
        )
        data = getattr(resp, "data", None)
        if not data or not getattr(data[0], "embedding", None):
            raise ValueError("No embedding data received")
        return data[0].embedding

    def embed_documents(self, texts: List[str]):
        resp = self.client.embeddings.create(
            model=self.model,
            input=texts,
            encoding_format="float",
            extra_headers=self.extra_headers or None,
        )
        data = getattr(resp, "data", None)
        if not data:
            raise ValueError("No embedding data received")
        return [item.embedding for item in data]


class DocumentProcessor:
    """Service for processing PDF documents and creating vector embeddings.
    
    This service handles loading PDF files, splitting them into chunks,
    creating embeddings, and storing them in a vector database.
    """
    
    def __init__(self):
        """Initialize the document processor."""
        self.docs_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "docs")
        self.persist_root = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "vectordb")
        
        os.makedirs(self.persist_root, exist_ok=True)
        
        openrouter_key = settings.OPENROUTER_API_KEY
        openai_key = settings.OPENAI_API_KEY
        try:
            if openrouter_key and openrouter_key != "":
                provider = "openrouter"
                model_name = "qwen/qwen3-embedding-4b"
                self.embeddings = OpenRouterEmbeddings(
                    api_key=openrouter_key,
                    base_url=getattr(settings, "OPENROUTER_API_BASE", "https://openrouter.ai/api/v1"),
                    model=model_name,
                )
                probe = self.embeddings.embed_query("probe")
                dim = len(probe) if isinstance(probe, (list, tuple)) else 0
                self.embeddings_available = True
                # logging.info("OpenRouter embeddings initialized: qwen/qwen3-embedding-0.6b")
                logging.info("OpenRouter embeddings initialized: qwen/qwen3-embedding-4b")
            elif openai_key and openai_key != "":
                provider = "openai"
                model_name = settings.OPENAI_EMBEDDING_MODEL
                self.embeddings = OpenAIEmbeddings(
                    openai_api_key=openai_key,
                    model=model_name,
                )
                probe = self.embeddings.embed_query("probe")
                dim = len(probe) if isinstance(probe, (list, tuple)) else 0
                self.embeddings_available = True
                logging.info("OpenAI embeddings initialized")
            else:
                self.embeddings_available = False
                self.embeddings = None
                logging.error("Embeddings are not configured. Provide OPENROUTER_API_KEY or OPENAI_API_KEY.")
                provider = "none"
                model_name = "none"
                dim = 0
        except Exception as e:
            logging.error(f"Error initializing embeddings: {str(e)}")
            self.embeddings_available = False
            self.embeddings = None
            provider = "error"
            model_name = "error"
            dim = 0

        self.persist_directory = os.path.join(
            self.persist_root,
            f"{provider}-{model_name.replace('/', '_')}-{dim}"
        )
        os.makedirs(self.persist_directory, exist_ok=True)

        
        # Initialize text splitter for chunking documents
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
        
        # Persian configuration parameters
        self.persian_config = {
            'chunk_size': 1500,
            'chunk_overlap': 200,
            'persian_separators': ["\n\n", "\n", ".", "،", "؛", " ", ""],
            'table_extraction': True,
            'save_intermediate': True,
            'add_start_index': True,
            'length_function': len,
            'output_format': 'markdown'
        }
        
        # Initialize Persian text splitter
        self.persian_text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.persian_config['chunk_size'],
            chunk_overlap=self.persian_config['chunk_overlap'],
            length_function=len,
            separators=self.persian_config['persian_separators']
        )
        
        # Initialize vector store
        self._vector_store = None
    
    def get_vector_store(self):
        """Get or create the vector store."""
        # If embeddings are not available, return None
        if not hasattr(self, 'embeddings_available') or not self.embeddings_available:
            logging.warning("Cannot access vector store: embeddings are not available.")
            return None
            
        if self._vector_store is None:
            # Check if vector store exists
            if os.path.exists(self.persist_directory) and len(os.listdir(self.persist_directory)) > 0:
                # Load existing vector store
                self._vector_store = Chroma(
                    persist_directory=self.persist_directory,
                    embedding_function=self.embeddings
                )
            else:
                # Create new vector store
                self._vector_store = Chroma(
                    persist_directory=self.persist_directory,
                    embedding_function=self.embeddings
                )
        
        return self._vector_store
    
    def process_pdf(self, file_path: str) -> List[Document]:
        """Process a single PDF file and return the extracted documents.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            List of Document objects
        """
        try:
            # Load PDF
            loader = PyPDFLoader(file_path)
            documents = loader.load()
            
            # Extract metadata from filename
            filename = os.path.basename(file_path)
            for doc in documents:
                doc.metadata["source"] = filename
                doc.metadata["file_path"] = file_path
            
            # Split documents into chunks
            split_docs = self.text_splitter.split_documents(documents)
            
            return split_docs
        except Exception as e:
            print(f"Error processing PDF {file_path}: {str(e)}")
            return []
    
    def process_all_pdfs(self) -> int:
        """Process all PDF files in the docs directory and add them to the vector store.
        
        Returns:
            Number of documents processed
        """
        all_docs = []
        
        # Get all PDF files in the docs directory
        pdf_files = [os.path.join(self.docs_dir, f) for f in os.listdir(self.docs_dir) 
                    if f.lower().endswith('.pdf')]
        
        # Process each PDF file
        for pdf_file in pdf_files:
            docs = self.process_pdf(pdf_file)
            all_docs.extend(docs)
        
        # Add documents to vector store in batches to avoid token limit issues
        if all_docs:
            vector_store = self.get_vector_store()
            
            # Process in batches of 100 documents to stay well under the 300k token limit
            batch_size = 100
            for i in range(0, len(all_docs), batch_size):
                batch = all_docs[i:i + batch_size]
                vector_store.add_documents(batch)
                print(f"Processed batch {i//batch_size + 1}/{(len(all_docs) + batch_size - 1)//batch_size} with {len(batch)} documents")
            
            vector_store.persist()
        
        return len(all_docs)
    
    def search_documents(self, query: str, k: int = 4) -> List[Dict[str, Any]]:
        """Search for documents relevant to the query.
        
        Args:
            query: The search query
            k: Number of documents to return
            
        Returns:
            List of documents with their content and metadata
        """
        vector_store = self.get_vector_store()
        docs = vector_store.similarity_search(query, k=k)
        
        results = []
        for doc in docs:
            results.append({
                "content": doc.page_content,
                "metadata": doc.metadata
            })
        
        return results
    
    def clean_persian_text(self, text: str) -> str:
        """Clean and normalize Persian text.
        
        Args:
            text: Raw Persian text
            
        Returns:
            Cleaned and normalized Persian text
        """
        if not text:
            return ""
        
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text.strip())
        
        # Convert Arabic numerals (٠-٩) to Persian (۰-۹)
        arabic_to_persian = {
            '٠': '۰', '١': '۱', '٢': '۲', '٣': '۳', '٤': '۴',
            '٥': '۵', '٦': '۶', '٧': '۷', '٨': '۸', '٩': '۹'
        }
        
        for arabic, persian in arabic_to_persian.items():
            text = text.replace(arabic, persian)
        
        # Remove control characters except newlines and tabs
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        
        # Normalize text encoding (remove zero-width characters)
        text = re.sub(r'[\u200B-\u200D\uFEFF]', '', text)
        
        return text
    
    def detect_rtl_headers(self, text: str) -> str:
        """Detect and format Persian headers in text.
        
        Args:
            text: Persian text with potential headers
            
        Returns:
            Text with Markdown-formatted headers
        """
        lines = text.split('\n')
        formatted_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                formatted_lines.append(line)
                continue
            
            # Detect headers based on patterns
            header_level = 0
            
            # Pattern 1: Lines ending with colon
            if line.endswith(':') and len(line) < 100:
                header_level = 2
            
            # Pattern 2: Numbered sections (e.g., "۱. عنوان")
            elif re.match(r'^[۰-۹]+[.)]\s+', line):
                header_level = 3
            
            # Pattern 3: Short lines that might be titles (less than 50 chars)
            elif len(line) < 50 and not line.endswith('.') and not line.endswith('،'):
                # Check if it contains mostly Persian characters
                persian_chars = len(re.findall(r'[\u0600-\u06FF]', line))
                if persian_chars > len(line) * 0.5:
                    header_level = 2
            
            # Apply header formatting
            if header_level > 0:
                line = '#' * header_level + ' ' + line
            
            formatted_lines.append(line)
        
        return '\n'.join(formatted_lines)
    
    def reshape_persian_for_display(self, text: str) -> str:
        """Reshape Persian text for proper RTL display.
        
        Args:
            text: Persian text
            
        Returns:
            Properly shaped text for RTL display
        """
        if not PERSIAN_SUPPORT or not text:
            return text
        
        try:
            # Reshape Arabic/Persian text for proper character joining
            reshaped_text = reshape(text)
            # Apply bidirectional algorithm for proper RTL display
            display_text = get_display(reshaped_text)
            return display_text
        except Exception as e:
            logging.warning(f"Error reshaping Persian text: {e}")
            return text
    
    def extract_tables_from_pdf(self, pdf_path: str) -> List[Dict[str, Any]]:
        """Extract tables from PDF using PyMuPDF.
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            List of dictionaries containing table data with page numbers
        """
        tables_data = []
        
        try:
            # Open PDF document
            doc = pymupdf.open(pdf_path)
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # Find tables on the page
                tables = page.find_tables()
                
                for table_index, table in enumerate(tables):
                    try:
                        # Extract table data
                        table_data = table.extract()
                        
                        if table_data and len(table_data) > 1:  # Ensure we have header and data rows
                            # Convert to pandas DataFrame
                            df = pd.DataFrame(table_data[1:], columns=table_data[0])
                            
                            # Clean Persian text in table cells
                            for col in df.columns:
                                df[col] = df[col].astype(str).apply(self.clean_persian_text)
                            
                            # Convert to Markdown format
                            markdown_table = df.to_markdown(index=False, tablefmt='pipe')
                            
                            # Store table information
                            table_info = {
                                'page_number': page_num + 1,
                                'table_index': table_index + 1,
                                'dataframe': df,
                                'markdown': markdown_table,
                                'bbox': table.bbox,  # Bounding box coordinates
                                'rows': len(df),
                                'columns': len(df.columns)
                            }
                            
                            tables_data.append(table_info)
                            
                    except Exception as e:
                        logging.warning(f"Error extracting table {table_index + 1} from page {page_num + 1}: {e}")
                        continue
            
            doc.close()
            
        except Exception as e:
            logging.error(f"Error processing PDF {pdf_path}: {e}")
        
        return tables_data
    
    def load_persian_pdf_as_markdown(self, pdf_path: str, output_dir: str = None) -> str:
        """Load Persian PDF and convert to Markdown format.
        
        Args:
            pdf_path: Path to the PDF file (relative to project root)
            output_dir: Output directory for saving files
            
        Returns:
            Markdown content as string
        """
        # Convert relative path to absolute
        if not os.path.isabs(pdf_path):
            pdf_path = os.path.join(os.getcwd(), pdf_path)
        
        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"PDF file not found: {pdf_path}")
        
        markdown_content = []
        filename = Path(pdf_path).stem
        
        try:
            # Extract text using PyMuPDF for better Persian support
            doc = pymupdf.open(pdf_path)
            
            # Add document title
            markdown_content.append(f"# {filename}\n")
            
            # Extract tables first
            tables = self.extract_tables_from_pdf(pdf_path) if self.persian_config['table_extraction'] else []
            tables_by_page = {table['page_number']: [] for table in tables}
            for table in tables:
                tables_by_page[table['page_number']].append(table)
            
            # Process each page
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # Add page header
                markdown_content.append(f"\n## صفحه {page_num + 1}\n")
                
                # Extract text from page
                text = page.get_text()
                
                # Clean and normalize Persian text
                cleaned_text = self.clean_persian_text(text)
                
                # Detect and format headers
                formatted_text = self.detect_rtl_headers(cleaned_text)
                
                # Add page content
                if formatted_text.strip():
                    markdown_content.append(formatted_text)
                
                # Add tables for this page
                if page_num + 1 in tables_by_page:
                    for table in tables_by_page[page_num + 1]:
                        markdown_content.append(f"\n### جدول {table['table_index']}\n")
                        markdown_content.append(table['markdown'])
                        markdown_content.append("\n")
            
            doc.close()
            
            # Join all content
            final_markdown = "\n".join(markdown_content)
            
            # Save to file if output directory is specified
            if output_dir and self.persian_config['save_intermediate']:
                output_path = Path(output_dir)
                markdown_dir = output_path / "markdown"
                markdown_dir.mkdir(parents=True, exist_ok=True)
                
                markdown_file = markdown_dir / f"{filename}.md"
                with open(markdown_file, 'w', encoding='utf-8') as f:
                    f.write(final_markdown)
                
                logging.info(f"Markdown saved to: {markdown_file}")
            
            return final_markdown
            
        except Exception as e:
            logging.error(f"Error converting PDF to Markdown: {e}")
            raise
    
    def process_persian_pdf_directory(self, input_dir: str, output_dir: str, create_vectors: bool = True) -> Dict[str, Any]:
        """Process all PDF files in a directory.
        
        Args:
            input_dir: Directory containing PDF files (relative to project root)
            output_dir: Output directory for processed files
            create_vectors: Whether to create vector embeddings
            
        Returns:
            Processing statistics dictionary
        """
        # Convert relative paths to absolute
        if not os.path.isabs(input_dir):
            input_dir = os.path.join(os.getcwd(), input_dir)
        if not os.path.isabs(output_dir):
            output_dir = os.path.join(os.getcwd(), output_dir)
        
        if not os.path.exists(input_dir):
            raise FileNotFoundError(f"Input directory not found: {input_dir}")
        
        # Create output directories
        output_path = Path(output_dir)
        markdown_dir = output_path / "markdown"
        json_dir = output_path / "structured_json"
        vector_dir = output_path / "vector_db"
        
        for dir_path in [markdown_dir, json_dir, vector_dir]:
            dir_path.mkdir(parents=True, exist_ok=True)
        
        # Initialize statistics
        stats = {
            'total_files': 0,
            'successful': 0,
            'failed': 0,
            'total_pages': 0,
            'total_tables': 0,
            'processing_time': 0,
            'start_time': datetime.now().isoformat(),
            'processed_files': [],
            'failed_files': []
        }
        
        start_time = datetime.now()
        
        # Find all PDF files
        pdf_files = list(Path(input_dir).glob("*.pdf"))
        stats['total_files'] = len(pdf_files)
        
        documents_for_vector = []
        
        for pdf_file in pdf_files:
            try:
                logging.info(f"Processing: {pdf_file.name}")
                
                # Convert PDF to Markdown
                markdown_content = self.load_persian_pdf_as_markdown(
                    str(pdf_file), 
                    output_dir
                )
                
                # Extract tables
                tables = self.extract_tables_from_pdf(str(pdf_file))
                
                # Generate file checksum
                file_hash = hashlib.md5(pdf_file.read_bytes()).hexdigest()
                
                # Create structured data
                structured_data = {
                    'filename': pdf_file.name,
                    'file_path': str(pdf_file),
                    'checksum': file_hash,
                    'processed_at': datetime.now().isoformat(),
                    'total_pages': 0,
                    'total_tables': len(tables),
                    'tables': tables,
                    'markdown_content': markdown_content
                }
                
                # Count pages
                try:
                    doc = pymupdf.open(str(pdf_file))
                    structured_data['total_pages'] = len(doc)
                    doc.close()
                except:
                    structured_data['total_pages'] = 0
                
                # Save structured JSON
                json_file = json_dir / f"{pdf_file.stem}.json"
                with open(json_file, 'w', encoding='utf-8') as f:
                    json.dump(structured_data, f, ensure_ascii=False, indent=2, default=str)
                
                # Prepare documents for vector store
                if create_vectors:
                    # Split text using Persian-aware splitter
                    chunks = self.persian_text_splitter.split_text(markdown_content)
                    
                    for i, chunk in enumerate(chunks):
                        # Check if chunk contains table
                        has_table = 'جدول' in chunk or '|' in chunk
                        
                        doc = Document(
                            page_content=chunk,
                            metadata={
                                'source': str(pdf_file),
                                'filename': pdf_file.name,
                                'chunk_index': i,
                                'has_table': has_table,
                                'file_checksum': file_hash,
                                'processed_at': structured_data['processed_at']
                            }
                        )
                        documents_for_vector.append(doc)
                
                # Update statistics
                stats['successful'] += 1
                stats['total_pages'] += structured_data['total_pages']
                stats['total_tables'] += len(tables)
                stats['processed_files'].append({
                    'filename': pdf_file.name,
                    'pages': structured_data['total_pages'],
                    'tables': len(tables),
                    'checksum': file_hash
                })
                
            except Exception as e:
                logging.error(f"Error processing {pdf_file.name}: {e}")
                stats['failed'] += 1
                stats['failed_files'].append({
                    'filename': pdf_file.name,
                    'error': str(e)
                })
        
        # Add documents to vector store
        if create_vectors and documents_for_vector:
            try:
                vector_store = self.get_vector_store()
                vector_store.add_documents(documents_for_vector)
                vector_store.persist()
                logging.info(f"Added {len(documents_for_vector)} document chunks to vector store")
            except Exception as e:
                logging.error(f"Error adding documents to vector store: {e}")
        
        # Calculate processing time
        end_time = datetime.now()
        stats['processing_time'] = (end_time - start_time).total_seconds()
        stats['end_time'] = end_time.isoformat()
        
        # Save statistics
        stats_file = output_path / "processing_stats.json"
        with open(stats_file, 'w', encoding='utf-8') as f:
            json.dump(stats, f, ensure_ascii=False, indent=2)
        
        logging.info(f"Processing complete. Stats saved to: {stats_file}")
        return stats
    
    def batch_process_directory(self, input_dir: str, output_dir: str, create_vectors: bool = True) -> Dict[str, Any]:
        """Batch process multiple PDFs with progress tracking.
        
        Args:
            input_dir: Directory containing PDF files
            output_dir: Output directory for processed files
            create_vectors: Whether to create vector embeddings
            
        Returns:
            Processing statistics with progress information
        """
        # Convert relative paths to absolute
        if not os.path.isabs(input_dir):
            input_dir = os.path.join(os.getcwd(), input_dir)
        if not os.path.isabs(output_dir):
            output_dir = os.path.join(os.getcwd(), output_dir)
        
        if not os.path.exists(input_dir):
            raise FileNotFoundError(f"Input directory not found: {input_dir}")
        
        # Create output directories
        output_path = Path(output_dir)
        for subdir in ["markdown", "structured_json", "vector_db"]:
            (output_path / subdir).mkdir(parents=True, exist_ok=True)
        
        # Find all PDF files
        pdf_files = list(Path(input_dir).glob("*.pdf"))
        total_files = len(pdf_files)
        
        if total_files == 0:
            logging.warning(f"No PDF files found in {input_dir}")
            return {'total_files': 0, 'successful': 0, 'failed': 0}
        
        # Initialize progress tracking
        stats = {
            'total_files': total_files,
            'successful': 0,
            'failed': 0,
            'total_pages': 0,
            'total_tables': 0,
            'processing_time': 0,
            'start_time': datetime.now().isoformat(),
            'processed_files': [],
            'failed_files': [],
            'progress_percentage': 0
        }
        
        start_time = datetime.now()
        documents_for_vector = []
        
        # Process files with progress tracking
        for i, pdf_file in enumerate(pdf_files, 1):
            try:
                # Update progress
                progress = (i / total_files) * 100
                stats['progress_percentage'] = round(progress, 2)
                
                logging.info(f"Processing {i}/{total_files} ({progress:.1f}%): {pdf_file.name}")
                
                # Process single PDF
                markdown_content = self.load_persian_pdf_as_markdown(
                    str(pdf_file), 
                    output_dir
                )
                
                # Extract tables and metadata
                tables = self.extract_tables_from_pdf(str(pdf_file))
                file_hash = hashlib.md5(pdf_file.read_bytes()).hexdigest()
                
                # Count pages
                page_count = 0
                try:
                    doc = pymupdf.open(str(pdf_file))
                    page_count = len(doc)
                    doc.close()
                except:
                    pass
                
                # Create structured data
                structured_data = {
                    'filename': pdf_file.name,
                    'file_path': str(pdf_file),
                    'checksum': file_hash,
                    'processed_at': datetime.now().isoformat(),
                    'total_pages': page_count,
                    'total_tables': len(tables),
                    'tables': [{
                        'page_number': table['page_number'],
                        'table_index': table['table_index'],
                        'rows': table['rows'],
                        'columns': table['columns'],
                        'markdown': table['markdown']
                    } for table in tables],
                    'markdown_content': markdown_content
                }
                
                # Save structured JSON
                json_file = output_path / "structured_json" / f"{pdf_file.stem}.json"
                with open(json_file, 'w', encoding='utf-8') as f:
                    json.dump(structured_data, f, ensure_ascii=False, indent=2, default=str)
                
                # Prepare for vector store
                if create_vectors:
                    chunks = self.persian_text_splitter.split_text(markdown_content)
                    
                    for chunk_idx, chunk in enumerate(chunks):
                        has_table = 'جدول' in chunk or '|' in chunk
                        
                        doc = Document(
                            page_content=chunk,
                            metadata={
                                'source': str(pdf_file),
                                'filename': pdf_file.name,
                                'chunk_index': chunk_idx,
                                'has_table': has_table,
                                'file_checksum': file_hash,
                                'total_pages': page_count,
                                'processed_at': structured_data['processed_at']
                            }
                        )
                        documents_for_vector.append(doc)
                
                # Update statistics
                stats['successful'] += 1
                stats['total_pages'] += page_count
                stats['total_tables'] += len(tables)
                stats['processed_files'].append({
                    'filename': pdf_file.name,
                    'pages': page_count,
                    'tables': len(tables),
                    'checksum': file_hash,
                    'processing_order': i
                })
                
            except Exception as e:
                logging.error(f"Error processing {pdf_file.name}: {e}")
                stats['failed'] += 1
                stats['failed_files'].append({
                    'filename': pdf_file.name,
                    'error': str(e),
                    'processing_order': i
                })
        
        # Add to vector store in batches
        if create_vectors and documents_for_vector:
            try:
                batch_size = 100
                vector_store = self.get_vector_store()
                
                for i in range(0, len(documents_for_vector), batch_size):
                    batch = documents_for_vector[i:i + batch_size]
                    vector_store.add_documents(batch)
                    logging.info(f"Added batch {i//batch_size + 1} ({len(batch)} documents) to vector store")
                
                vector_store.persist()
                logging.info(f"Successfully added {len(documents_for_vector)} document chunks to vector store")
                
            except Exception as e:
                logging.error(f"Error adding documents to vector store: {e}")
        
        # Finalize statistics
        end_time = datetime.now()
        stats['processing_time'] = (end_time - start_time).total_seconds()
        stats['end_time'] = end_time.isoformat()
        stats['progress_percentage'] = 100.0
        
        # Calculate averages
        if stats['successful'] > 0:
            stats['avg_pages_per_file'] = round(stats['total_pages'] / stats['successful'], 2)
            stats['avg_tables_per_file'] = round(stats['total_tables'] / stats['successful'], 2)
            stats['avg_processing_time_per_file'] = round(stats['processing_time'] / stats['successful'], 2)
        
        # Save comprehensive statistics
        stats_file = output_path / "batch_processing_stats.json"
        with open(stats_file, 'w', encoding='utf-8') as f:
            json.dump(stats, f, ensure_ascii=False, indent=2)
        
        logging.info(f"Batch processing complete. Processed {stats['successful']}/{stats['total_files']} files successfully")
        logging.info(f"Total pages: {stats['total_pages']}, Total tables: {stats['total_tables']}")
        logging.info(f"Processing time: {stats['processing_time']:.2f} seconds")
        logging.info(f"Statistics saved to: {stats_file}")
        
        return stats
    
    def extract_tables_from_docx(self, docx_path: str) -> List[Dict[str, Any]]:
        """Extract tables from DOCX file.
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            List of dictionaries containing table data
        """
        if not DOCX_SUPPORT:
            logging.error("DOCX support not available. Install python-docx library.")
            return []
        
        tables_data = []
        
        try:
            doc = DocxDocument(docx_path)
            
            for table_index, table in enumerate(doc.tables):
                try:
                    # Extract table data
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        table_data.append(row_data)
                    
                    if table_data and len(table_data) > 1:  # Ensure we have header and data rows
                        # Convert to pandas DataFrame
                        df = pd.DataFrame(table_data[1:], columns=table_data[0])
                        
                        # Clean Persian text in table cells
                        for col in df.columns:
                            df[col] = df[col].astype(str).apply(self.clean_persian_text)
                        
                        # Convert to Markdown format
                        markdown_table = df.to_markdown(index=False, tablefmt='pipe')
                        
                        # Store table information
                        table_info = {
                            'table_index': table_index + 1,
                            'dataframe': df,
                            'markdown': markdown_table,
                            'rows': len(df),
                            'columns': len(df.columns)
                        }
                        
                        tables_data.append(table_info)
                        
                except Exception as e:
                    logging.warning(f"Error extracting table {table_index + 1}: {e}")
                    continue
            
        except Exception as e:
            logging.error(f"Error processing DOCX {docx_path}: {e}")
        
        return tables_data
    
    def load_docx_as_markdown(self, docx_path: str, output_dir: str = None) -> str:
        """Load DOCX file and convert to Markdown format.
        
        Args:
            docx_path: Path to the DOCX file (relative to project root)
            output_dir: Output directory for saving files
            
        Returns:
            Markdown content as string
        """
        if not DOCX_SUPPORT:
            raise ImportError("DOCX support not available. Install python-docx library.")
        
        # Convert relative path to absolute
        if not os.path.isabs(docx_path):
            docx_path = os.path.join(os.getcwd(), docx_path)
        
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"DOCX file not found: {docx_path}")
        
        markdown_content = []
        filename = Path(docx_path).stem
        
        try:
            # Load DOCX document
            doc = DocxDocument(docx_path)
            
            # Add document title
            markdown_content.append(f"# {filename}\n")
            
            # Extract tables first
            tables = self.extract_tables_from_docx(docx_path)
            table_index = 0
            
            # Process paragraphs and tables in order
            for element in doc.element.body:
                # Check if element is a paragraph
                if element.tag.endswith('p'):
                    # Find the corresponding paragraph object
                    for para in doc.paragraphs:
                        if para._element == element:
                            text = para.text.strip()
                            if text:
                                # Clean Persian text
                                cleaned_text = self.clean_persian_text(text)
                                
                                # Check if it's a heading based on style
                                if para.style.name.startswith('Heading'):
                                    heading_level = int(para.style.name.replace('Heading ', ''))
                                    markdown_content.append(f"\n{'#' * heading_level} {cleaned_text}\n")
                                else:
                                    # Detect headers using existing logic
                                    formatted_text = self.detect_rtl_headers(cleaned_text)
                                    markdown_content.append(formatted_text)
                            break
                
                # Check if element is a table
                elif element.tag.endswith('tbl'):
                    if table_index < len(tables):
                        table = tables[table_index]
                        markdown_content.append(f"\n### جدول {table['table_index']}\n")
                        markdown_content.append(table['markdown'])
                        markdown_content.append("\n")
                        table_index += 1
            
            # Join all content
            final_markdown = "\n".join(markdown_content)
            
            # Save to file if output directory is specified
            if output_dir and self.persian_config['save_intermediate']:
                output_path = Path(output_dir)
                markdown_dir = output_path / "markdown"
                markdown_dir.mkdir(parents=True, exist_ok=True)
                
                markdown_file = markdown_dir / f"{filename}.md"
                with open(markdown_file, 'w', encoding='utf-8') as f:
                    f.write(final_markdown)
                
                logging.info(f"Markdown saved to: {markdown_file}")
            
            return final_markdown
            
        except Exception as e:
            logging.error(f"Error converting DOCX to Markdown: {e}")
            raise
    
    def process_docx(self, file_path: str) -> List[Document]:
        """Process a single DOCX file and return the extracted documents.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            List of Document objects
        """
        if not DOCX_SUPPORT:
            logging.error("DOCX support not available. Install python-docx library.")
            return []
        
        try:
            # Load DOCX and convert to markdown
            markdown_content = self.load_docx_as_markdown(file_path)
            
            # Extract metadata from filename
            filename = os.path.basename(file_path)
            
            # Split content into chunks
            chunks = self.persian_text_splitter.split_text(markdown_content)
            
            # Create Document objects
            documents = []
            for i, chunk in enumerate(chunks):
                doc = Document(
                    page_content=chunk,
                    metadata={
                        "source": filename,
                        "file_path": file_path,
                        "chunk_index": i,
                        "file_type": "docx"
                    }
                )
                documents.append(doc)
            
            return documents
            
        except Exception as e:
            logging.error(f"Error processing DOCX {file_path}: {str(e)}")
            return []
    
    def process_all_docx(self) -> int:
        """Process all DOCX files in the docs directory and add them to the vector store.
        
        Returns:
            Number of documents processed
        """
        if not DOCX_SUPPORT:
            logging.error("DOCX support not available. Install python-docx library.")
            return 0
        
        all_docs = []
        
        # Get all DOCX files in the docs directory
        docx_files = [os.path.join(self.docs_dir, f) for f in os.listdir(self.docs_dir) 
                     if f.lower().endswith(('.docx', '.doc'))]
        
        # Process each DOCX file
        for docx_file in docx_files:
            docs = self.process_docx(docx_file)
            all_docs.extend(docs)
        
        # Add documents to vector store in batches
        if all_docs:
            vector_store = self.get_vector_store()
            
            # Process in batches of 100 documents
            batch_size = 100
            for i in range(0, len(all_docs), batch_size):
                batch = all_docs[i:i + batch_size]
                vector_store.add_documents(batch)
                logging.info(f"Processed batch {i//batch_size + 1}/{(len(all_docs) + batch_size - 1)//batch_size} with {len(batch)} documents")
            
            vector_store.persist()
        
        return len(all_docs)
    
    def process_all_documents(self) -> Dict[str, int]:
        """Process all PDF and DOCX files in the docs directory.
        
        Returns:
            Dictionary with counts of processed documents by type
        """
        pdf_count = self.process_all_pdfs()
        docx_count = self.process_all_docx()
        
        return {
            'pdf_documents': pdf_count,
            'docx_documents': docx_count,
            'total_documents': pdf_count + docx_count
        }
    
    def batch_process_mixed_directory(self, input_dir: str, output_dir: str, create_vectors: bool = True) -> Dict[str, Any]:
        """Batch process both PDF and DOCX files with progress tracking.
        
        Args:
            input_dir: Directory containing PDF and DOCX files
            output_dir: Output directory for processed files
            create_vectors: Whether to create vector embeddings
            
        Returns:
            Processing statistics with progress information
        """
        # Convert relative paths to absolute
        if not os.path.isabs(input_dir):
            input_dir = os.path.join(os.getcwd(), input_dir)
        if not os.path.isabs(output_dir):
            output_dir = os.path.join(os.getcwd(), output_dir)
        
        if not os.path.exists(input_dir):
            raise FileNotFoundError(f"Input directory not found: {input_dir}")
        
        # Create output directories
        output_path = Path(output_dir)
        for subdir in ["markdown", "structured_json", "vector_db"]:
            (output_path / subdir).mkdir(parents=True, exist_ok=True)
        
        # Find all PDF and DOCX files
        pdf_files = list(Path(input_dir).glob("*.pdf"))
        docx_files = list(Path(input_dir).glob("*.docx"))
        all_files = pdf_files + docx_files
        total_files = len(all_files)
        
        if total_files == 0:
            logging.warning(f"No PDF or DOCX files found in {input_dir}")
            return {'total_files': 0, 'successful': 0, 'failed': 0}
        
        # Initialize progress tracking
        stats = {
            'total_files': total_files,
            'pdf_files': len(pdf_files),
            'docx_files': len(docx_files),
            'successful': 0,
            'failed': 0,
            'total_pages': 0,
            'total_tables': 0,
            'processing_time': 0,
            'start_time': datetime.now().isoformat(),
            'processed_files': [],
            'failed_files': [],
            'progress_percentage': 0
        }
        
        start_time = datetime.now()
        documents_for_vector = []
        
        # Process files with progress tracking
        for i, file_path in enumerate(all_files, 1):
            try:
                # Update progress
                progress = (i / total_files) * 100
                stats['progress_percentage'] = round(progress, 2)
                
                file_type = 'pdf' if file_path.suffix.lower() == '.pdf' else 'docx'
                logging.info(f"Processing {i}/{total_files} ({progress:.1f}%): {file_path.name} [{file_type.upper()}]")
                
                # Process based on file type
                if file_type == 'pdf':
                    markdown_content = self.load_persian_pdf_as_markdown(str(file_path), output_dir)
                    tables = self.extract_tables_from_pdf(str(file_path))
                    
                    # Count pages for PDF
                    page_count = 0
                    try:
                        doc = pymupdf.open(str(file_path))
                        page_count = len(doc)
                        doc.close()
                    except:
                        pass
                else:  # docx
                    if not DOCX_SUPPORT:
                        logging.error(f"Skipping {file_path.name}: DOCX support not available")
                        stats['failed'] += 1
                        continue
                    
                    markdown_content = self.load_docx_as_markdown(str(file_path), output_dir)
                    tables = self.extract_tables_from_docx(str(file_path))
                    page_count = 1  # DOCX doesn't have fixed pages
                
                # Generate file checksum
                file_hash = hashlib.md5(file_path.read_bytes()).hexdigest()
                
                # Create structured data
                structured_data = {
                    'filename': file_path.name,
                    'file_path': str(file_path),
                    'file_type': file_type,
                    'checksum': file_hash,
                    'processed_at': datetime.now().isoformat(),
                    'total_pages': page_count,
                    'total_tables': len(tables),
                    'tables': [{
                        'table_index': table['table_index'],
                        'rows': table['rows'],
                        'columns': table['columns'],
                        'markdown': table['markdown']
                    } for table in tables],
                    'markdown_content': markdown_content
                }
                
                # Add page number for PDF tables
                if file_type == 'pdf':
                    for idx, table in enumerate(tables):
                        if 'page_number' in table:
                            structured_data['tables'][idx]['page_number'] = table['page_number']
                
                # Save structured JSON
                json_file = output_path / "structured_json" / f"{file_path.stem}.json"
                with open(json_file, 'w', encoding='utf-8') as f:
                    json.dump(structured_data, f, ensure_ascii=False, indent=2, default=str)
                
                # Prepare for vector store
                if create_vectors:
                    chunks = self.persian_text_splitter.split_text(markdown_content)
                    
                    for chunk_idx, chunk in enumerate(chunks):
                        has_table = 'جدول' in chunk or '|' in chunk
                        
                        doc = Document(
                            page_content=chunk,
                            metadata={
                                'source': str(file_path),
                                'filename': file_path.name,
                                'file_type': file_type,
                                'chunk_index': chunk_idx,
                                'has_table': has_table,
                                'file_checksum': file_hash,
                                'total_pages': page_count,
                                'processed_at': structured_data['processed_at']
                            }
                        )
                        documents_for_vector.append(doc)
                
                # Update statistics
                stats['successful'] += 1
                stats['total_pages'] += page_count
                stats['total_tables'] += len(tables)
                stats['processed_files'].append({
                    'filename': file_path.name,
                    'file_type': file_type,
                    'pages': page_count,
                    'tables': len(tables),
                    'checksum': file_hash,
                    'processing_order': i
                })
                
            except Exception as e:
                logging.error(f"Error processing {file_path.name}: {e}")
                stats['failed'] += 1
                stats['failed_files'].append({
                    'filename': file_path.name,
                    'file_type': file_type if 'file_type' in locals() else 'unknown',
                    'error': str(e),
                    'processing_order': i
                })
        
        # Add to vector store in batches
        if create_vectors and documents_for_vector:
            try:
                batch_size = 100
                vector_store = self.get_vector_store()
                
                for i in range(0, len(documents_for_vector), batch_size):
                    batch = documents_for_vector[i:i + batch_size]
                    vector_store.add_documents(batch)
                    logging.info(f"Added batch {i//batch_size + 1} ({len(batch)} documents) to vector store")
                
                vector_store.persist()
                logging.info(f"Successfully added {len(documents_for_vector)} document chunks to vector store")
                
            except Exception as e:
                logging.error(f"Error adding documents to vector store: {e}")
        
        # Finalize statistics
        end_time = datetime.now()
        stats['processing_time'] = (end_time - start_time).total_seconds()
        stats['end_time'] = end_time.isoformat()
        stats['progress_percentage'] = 100.0
        
        # Calculate averages
        if stats['successful'] > 0:
            stats['avg_pages_per_file'] = round(stats['total_pages'] / stats['successful'], 2)
            stats['avg_tables_per_file'] = round(stats['total_tables'] / stats['successful'], 2)
            stats['avg_processing_time_per_file'] = round(stats['processing_time'] / stats['successful'], 2)
        
        # Save comprehensive statistics
        stats_file = output_path / "batch_processing_stats.json"
        with open(stats_file, 'w', encoding='utf-8') as f:
            json.dump(stats, f, ensure_ascii=False, indent=2)
        
        logging.info(f"Batch processing complete. Processed {stats['successful']}/{stats['total_files']} files successfully")
        logging.info(f"PDF files: {stats['pdf_files']}, DOCX files: {stats['docx_files']}")
        logging.info(f"Total pages: {stats['total_pages']}, Total tables: {stats['total_tables']}")
        logging.info(f"Processing time: {stats['processing_time']:.2f} seconds")
        logging.info(f"Statistics saved to: {stats_file}")
        
        return stats

    async def sync_mongodb_to_vectordb(self, knowledge_base_service, batch_size: int = 10) -> Dict[str, Any]:
        """
        Sync user_contribution documents from MongoDB to vector database.
        
        This function retrieves documents with entry_type: "user_contribution" from MongoDB
        and adds them to the vector database using the add_knowledge_contribution function.
        
        Args:
            knowledge_base_service: An instance of KnowledgeBaseService
            batch_size: Number of documents to process in each batch
            
        Returns:
            Dictionary with sync statistics including processed_count, skipped_count, and error_count
        """
        stats = {
            'processed_count': 0,
            'skipped_count': 0,
            'error_count': 0,
            'start_time': datetime.now().isoformat(),
            'processed_documents': [],
            'failed_documents': []
        }
        
        try:
            # Initialize database service
            db_service = DatabaseService()
            await db_service.connect()
            
            
            
            # Retrieve user_contribution documents from MongoDB
            collection = db_service.get_knowledgebase_collection()
            
            # Query for user_contribution documents with synced filter
            query = {"entry_type": {"$in": ["user_contribution"]}, "synced": True}
            cursor = collection.find(query)
            documents = await cursor.to_list(length=None)
            
            logging.info(f"Found {len(documents)} user_contribution documents in MongoDB")
            
            # Process documents in batches
            for i in range(0, len(documents), batch_size):
                batch = documents[i:i + batch_size]
                
                for doc in batch:
                    try:
                        # Extract document fields
                        title = doc.get('title', '')
                        content = doc.get('content', '')
                        meta_tags = doc.get('meta_tags', [])
                        source = doc.get('source')
                        author_name = doc.get('author_name')
                        additional_references = doc.get('additional_references')
                        uploaded_file_path = doc.get('uploaded_file_path')
                        is_public = doc.get("is_public", False)
                        
                        # Skip documents with missing required fields
                        if not title or not content:
                            logging.warning(f"Skipping document {doc.get('_id')}: missing title or content")
                            stats['skipped_count'] += 1
                            continue
                        
                        # Add to vector database using existing function
                        result = await knowledge_base_service.add_knowledge_contribution(
                            title=title,
                            content=content,
                            meta_tags=meta_tags,
                            source=source,
                            author_name=author_name,
                            additional_references=additional_references,
                            uploaded_file_path=uploaded_file_path,
                            is_public=is_public
                        )
                        
                        stats['processed_count'] += 1
                        stats['processed_documents'].append({
                            'document_id': str(doc.get('_id')),
                            'title': title,
                            'result': result
                        })
                        
                        logging.info(f"Successfully processed document: {title}")
                        
                    except Exception as e:
                        logging.error(f"Error processing document {doc.get('_id')}: {str(e)}")
                        stats['error_count'] += 1
                        stats['failed_documents'].append({
                            'document_id': str(doc.get('_id')),
                            'title': doc.get('title', 'Unknown'),
                            'error': str(e)
                        })
                
                # Small delay between batches to avoid overwhelming the system
                import asyncio
                await asyncio.sleep(0.1)
            
            # Disconnect from database
            await db_service.disconnect()
            
            # Final statistics
            stats['end_time'] = datetime.now().isoformat()
            total_time = (datetime.fromisoformat(stats['end_time']) - 
                         datetime.fromisoformat(stats['start_time'])).total_seconds()
            stats['total_time_seconds'] = total_time
            
            logging.info(f"Sync completed: {stats['processed_count']} processed, "
                       f"{stats['skipped_count']} skipped, {stats['error_count']} errors")
            
        except Exception as e:
            logging.error(f"Failed to sync MongoDB to vector database: {str(e)}")
            stats['error_count'] += 1
            stats['error_message'] = str(e)
        
        return stats


# Singleton instance
_document_processor = None


def get_document_processor() -> DocumentProcessor:
    """Get the document processor instance.
    
    Returns:
        A singleton instance of the DocumentProcessor
    """
    global _document_processor
    if _document_processor is None:
        _document_processor = DocumentProcessor()
    return _document_processor